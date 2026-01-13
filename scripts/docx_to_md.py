#!/usr/bin/env python3
"""
Convert Word (.docx) documents to Markdown with rich formatting preservation.

Preserves:
- Headings (H1-H6)
- Bold, italic, underline, strikethrough
- Text colors (as HTML spans)
- Highlight/background colors
- Ordered and unordered lists (with nesting)
- Tables
- Hyperlinks
- Images (extracted to assets folder)
- Paragraphs and line breaks
- Blockquotes
- Code blocks (based on font)

Usage:
    python docx_to_md.py input.docx [output.md]
    python docx_to_md.py directory/  # Convert all .docx files
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


class DocxToMarkdown:
    """Convert a Word document to Markdown with formatting preservation."""

    # Namespaces used in docx XML
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    }

    def __init__(self, docx_path: str, output_path: Optional[str] = None):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)

        if output_path:
            self.output_path = Path(output_path)
        else:
            self.output_path = self.docx_path.with_suffix('.md')

        self.assets_dir = self.output_path.parent / f"{self.output_path.stem}_assets"
        self.image_counter = 0
        self.list_state = []  # Track nested list levels
        self.prev_was_list = False

    def convert(self) -> str:
        """Convert the document and return markdown string."""
        lines = []

        for element in self.doc.element.body:
            tag = element.tag.split('}')[-1]  # Remove namespace

            if tag == 'p':
                para = self._find_paragraph(element)
                if para:
                    md = self._convert_paragraph(para, element)
                    if md is not None:
                        lines.append(md)
            elif tag == 'tbl':
                table = self._find_table(element)
                if table:
                    lines.append(self._convert_table(table))
                    lines.append('')
            elif tag == 'sectPr':
                # Section properties - skip
                pass

        # Post-process: clean up excessive blank lines
        content = '\n'.join(lines)
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content.strip() + '\n'

    def _find_paragraph(self, element):
        """Find the python-docx Paragraph object for an XML element."""
        for para in self.doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, element):
        """Find the python-docx Table object for an XML element."""
        for table in self.doc.tables:
            if table._element is element:
                return table
        return None

    def _convert_paragraph(self, para, element) -> Optional[str]:
        """Convert a paragraph to markdown."""
        # Check for images first
        images = self._extract_images(element)

        # Get the text content with formatting
        text = self._get_formatted_text(para)

        # Handle empty paragraphs
        if not text.strip() and not images:
            self.prev_was_list = False
            return ''

        style_name = para.style.name if para.style else ''

        # Check for headings
        if style_name.startswith('Heading'):
            level = self._get_heading_level(style_name)
            self.prev_was_list = False
            self.list_state = []
            return f"{'#' * level} {text.strip()}"

        if style_name == 'Title':
            self.prev_was_list = False
            self.list_state = []
            return f"# {text.strip()}"

        if style_name == 'Subtitle':
            self.prev_was_list = False
            self.list_state = []
            return f"## {text.strip()}"

        # Check for list items
        list_info = self._get_list_info(element)
        if list_info:
            self.prev_was_list = True
            return self._format_list_item(text, list_info)

        # Check for blockquote (Quote style or indented)
        if 'Quote' in style_name or 'Block' in style_name:
            self.prev_was_list = False
            self.list_state = []
            return f"> {text.strip()}"

        # Add blank line after list if this isn't a list item
        prefix = ''
        if self.prev_was_list:
            prefix = '\n'
            self.list_state = []
        self.prev_was_list = False

        # Regular paragraph
        result = prefix + text.strip()

        # Add images
        if images:
            result += '\n' + '\n'.join(images)

        return result

    def _get_formatted_text(self, para) -> str:
        """Extract text from paragraph with inline formatting."""
        parts = []

        for run in para.runs:
            text = run.text
            if not text:
                continue

            # Get formatting
            bold = run.bold
            italic = run.italic
            underline = run.underline
            strike = run.font.strike

            # Get color
            color = self._get_run_color(run)
            highlight = self._get_highlight_color(run)

            # Check for code font (monospace)
            is_code = self._is_code_font(run)

            # Check for hyperlink
            hyperlink = self._get_hyperlink(run)

            # Apply formatting (inside out)
            formatted = text

            # Code formatting
            if is_code and not hyperlink:
                formatted = f"`{formatted}`"

            # Basic formatting
            if strike:
                formatted = f"~~{formatted}~~"
            if bold:
                formatted = f"**{formatted}**"
            if italic:
                formatted = f"*{formatted}*"

            # Underline (use HTML since MD doesn't support it)
            if underline:
                formatted = f"<u>{formatted}</u>"

            # Color (use HTML span)
            if color and color != '000000':
                formatted = f'<span style="color: #{color}">{formatted}</span>'

            # Highlight
            if highlight:
                formatted = f'<mark style="background-color: {highlight}">{formatted}</mark>'

            # Hyperlink
            if hyperlink:
                formatted = f"[{formatted}]({hyperlink})"

            parts.append(formatted)

        return ''.join(parts)

    def _get_run_color(self, run) -> Optional[str]:
        """Get the text color of a run as hex string."""
        try:
            color = run.font.color.rgb
            if color:
                return f"{color.red:02x}{color.green:02x}{color.blue:02x}"
        except (AttributeError, TypeError):
            pass

        # Check XML directly for theme colors
        try:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                color_elem = rPr.find(qn('w:color'))
                if color_elem is not None:
                    val = color_elem.get(qn('w:val'))
                    if val and val != 'auto':
                        return val
        except Exception:
            pass

        return None

    def _get_highlight_color(self, run) -> Optional[str]:
        """Get highlight/background color."""
        try:
            highlight = run.font.highlight_color
            if highlight:
                # Convert WD_COLOR_INDEX to CSS color
                color_map = {
                    1: '#FFFF00',   # Yellow
                    2: '#00FF00',   # Bright Green
                    3: '#00FFFF',   # Turquoise
                    4: '#FF00FF',   # Pink
                    5: '#0000FF',   # Blue
                    6: '#FF0000',   # Red
                    7: '#000080',   # Dark Blue
                    8: '#008080',   # Teal
                    9: '#008000',   # Green
                    10: '#800080',  # Violet
                    11: '#800000',  # Dark Red
                    12: '#808000',  # Dark Yellow
                    13: '#808080',  # Gray 50%
                    14: '#C0C0C0',  # Gray 25%
                    15: '#000000',  # Black
                }
                return color_map.get(highlight.real, None)
        except (AttributeError, TypeError):
            pass
        return None

    def _is_code_font(self, run) -> bool:
        """Check if run uses a monospace/code font."""
        try:
            font_name = run.font.name
            if font_name:
                code_fonts = ['Courier', 'Consolas', 'Monaco', 'Menlo',
                             'Source Code', 'Fira Code', 'monospace']
                return any(cf.lower() in font_name.lower() for cf in code_fonts)
        except (AttributeError, TypeError):
            pass
        return False

    def _get_hyperlink(self, run) -> Optional[str]:
        """Get hyperlink URL if run is part of a hyperlink."""
        try:
            # Check if parent is a hyperlink
            parent = run._element.getparent()
            if parent is not None and parent.tag.endswith('hyperlink'):
                rel_id = parent.get(qn('r:id'))
                if rel_id:
                    rel = self.doc.part.rels.get(rel_id)
                    if rel:
                        return rel.target_ref
        except Exception:
            pass
        return None

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r'(\d+)', style_name)
        if match:
            level = int(match.group(1))
            return min(level, 6)  # Max H6
        return 1

    def _get_list_info(self, element) -> Optional[dict]:
        """Get list information (type, level) from paragraph element."""
        try:
            pPr = element.find(qn('w:pPr'))
            if pPr is None:
                return None

            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                return None

            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))

            if ilvl is None or numId is None:
                return None

            level = int(ilvl.get(qn('w:val'), 0))
            num_id = int(numId.get(qn('w:val'), 0))

            # Determine if ordered or unordered by checking numbering definitions
            is_ordered = self._is_ordered_list(num_id, level)

            return {
                'level': level,
                'num_id': num_id,
                'ordered': is_ordered
            }
        except Exception:
            return None

    def _is_ordered_list(self, num_id: int, level: int) -> bool:
        """Check if a list is ordered based on numbering definition."""
        try:
            numbering = self.doc.part.numbering_part.numbering_definitions._numbering

            for num in numbering.findall(qn('w:num')):
                if int(num.get(qn('w:numId'), -1)) == num_id:
                    abstract_id = num.find(qn('w:abstractNumId'))
                    if abstract_id is not None:
                        abs_id = int(abstract_id.get(qn('w:val'), -1))

                        for abstract in numbering.findall(qn('w:abstractNum')):
                            if int(abstract.get(qn('w:abstractNumId'), -1)) == abs_id:
                                for lvl in abstract.findall(qn('w:lvl')):
                                    if int(lvl.get(qn('w:ilvl'), -1)) == level:
                                        numFmt = lvl.find(qn('w:numFmt'))
                                        if numFmt is not None:
                                            fmt = numFmt.get(qn('w:val'), '')
                                            return fmt in ['decimal', 'upperRoman',
                                                          'lowerRoman', 'upperLetter',
                                                          'lowerLetter']
        except Exception:
            pass
        return False

    def _format_list_item(self, text: str, list_info: dict) -> str:
        """Format a list item with proper indentation."""
        level = list_info['level']
        indent = '    ' * level

        if list_info['ordered']:
            # For ordered lists, we use 1. since MD renderers auto-number
            marker = '1.'
        else:
            marker = '-'

        return f"{indent}{marker} {text.strip()}"

    def _extract_images(self, element) -> list:
        """Extract images from paragraph and save to assets folder."""
        images = []

        try:
            # Find drawing elements
            for drawing in element.iter(qn('w:drawing')):
                for blip in drawing.iter(qn('a:blip')):
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        image_part = self.doc.part.related_parts.get(embed)
                        if image_part:
                            # Create assets directory if needed
                            self.assets_dir.mkdir(parents=True, exist_ok=True)

                            # Save image
                            self.image_counter += 1
                            ext = image_part.content_type.split('/')[-1]
                            if ext == 'jpeg':
                                ext = 'jpg'

                            image_name = f"image_{self.image_counter}.{ext}"
                            image_path = self.assets_dir / image_name

                            with open(image_path, 'wb') as f:
                                f.write(image_part.blob)

                            # Get alt text if available
                            alt_text = self._get_image_alt_text(drawing)

                            # Create relative path for markdown
                            rel_path = f"{self.assets_dir.name}/{image_name}"
                            images.append(f"![{alt_text}]({rel_path})")
        except Exception as e:
            pass

        return images

    def _get_image_alt_text(self, drawing) -> str:
        """Extract alt text from drawing element."""
        try:
            for docPr in drawing.iter(qn('wp:docPr')):
                name = docPr.get('name', '')
                descr = docPr.get('descr', '')
                return descr or name or 'image'
        except Exception:
            pass
        return 'image'

    def _convert_table(self, table) -> str:
        """Convert a table to markdown format."""
        if not table.rows:
            return ''

        lines = []

        # Process each row
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                # Get cell text with formatting
                cell_text = []
                for para in cell.paragraphs:
                    para_text = self._get_formatted_text(para)
                    if para_text.strip():
                        cell_text.append(para_text.strip())

                # Join multiple paragraphs with <br>
                cell_content = '<br>'.join(cell_text) if cell_text else ''
                # Escape pipes in cell content
                cell_content = cell_content.replace('|', '\\|')
                cells.append(cell_content)

            lines.append('| ' + ' | '.join(cells) + ' |')

            # Add header separator after first row
            if i == 0:
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')

        return '\n'.join(lines)

    def save(self):
        """Convert and save the markdown file."""
        content = self.convert()

        with open(self.output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return self.output_path


def convert_file(input_path: str, output_path: Optional[str] = None) -> Path:
    """Convert a single docx file to markdown."""
    converter = DocxToMarkdown(input_path, output_path)
    return converter.save()


def convert_directory(dir_path: str, recursive: bool = False) -> list:
    """Convert all docx files in a directory."""
    dir_path = Path(dir_path)
    pattern = '**/*.docx' if recursive else '*.docx'

    converted = []
    for docx_file in dir_path.glob(pattern):
        # Skip temporary files
        if docx_file.name.startswith('~$'):
            continue

        try:
            output = convert_file(str(docx_file))
            converted.append(output)
            print(f"✓ Converted: {docx_file.name} → {output.name}")
        except Exception as e:
            print(f"✗ Failed: {docx_file.name} - {e}")

    return converted


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word documents to Markdown with formatting preservation.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s document.docx                    # Convert single file
  %(prog)s document.docx output.md          # Convert with custom output name
  %(prog)s ./documents/                     # Convert all .docx in directory
  %(prog)s ./documents/ -r                  # Convert recursively
        """
    )

    parser.add_argument('input', help='Input .docx file or directory')
    parser.add_argument('output', nargs='?', help='Output .md file (optional)')
    parser.add_argument('-r', '--recursive', action='store_true',
                        help='Recursively convert files in subdirectories')

    args = parser.parse_args()

    input_path = Path(args.input)

    if not input_path.exists():
        print(f"Error: {args.input} does not exist")
        sys.exit(1)

    if input_path.is_dir():
        converted = convert_directory(str(input_path), args.recursive)
        print(f"\nConverted {len(converted)} file(s)")
    else:
        if not input_path.suffix.lower() == '.docx':
            print(f"Error: {args.input} is not a .docx file")
            sys.exit(1)

        output = convert_file(str(input_path), args.output)
        print(f"✓ Converted: {input_path.name} → {output}")


if __name__ == '__main__':
    main()
